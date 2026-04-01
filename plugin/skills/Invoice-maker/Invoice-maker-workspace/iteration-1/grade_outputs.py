#!/usr/bin/env python3
"""Grade invoice outputs against assertions in eval_metadata.json."""

import json
import sys
import glob
import os
from pathlib import Path

try:
    from docx import Document
except ImportError:
    os.system(f'{sys.executable} -m pip install python-docx --quiet --break-system-packages 2>/dev/null')
    from docx import Document


def extract_docx_text(docx_path):
    """Extract all text from a .docx file (paragraphs + tables)."""
    doc = Document(docx_path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return '\n'.join(texts)


def count_line_items(docx_path):
    """Count non-empty line item rows in the items table."""
    doc = Document(docx_path)
    if len(doc.tables) < 2:
        return 0
    t1 = doc.tables[1]
    count = 0
    for ri in range(1, len(t1.rows)):
        # A real line item row has "10%" in the GST column (col 4)
        gst_text = t1.rows[ri].cells[4].text.strip()
        if gst_text == '10%':
            count += 1
    return count


def check_wordsmithed(docx_path, raw_inputs):
    """Check that descriptions are not verbatim raw user input."""
    doc = Document(docx_path)
    if len(doc.tables) < 2:
        return False, "No items table found"
    t1 = doc.tables[1]
    for ri in range(1, len(t1.rows)):
        desc = t1.rows[ri].cells[2].text.strip() if len(t1.rows[ri].cells) > 2 else ''
        if not desc:
            continue
        for raw in raw_inputs:
            if raw.lower() == desc.lower():
                return False, f"Description '{desc}' matches raw input '{raw}' verbatim"
    return True, "Descriptions appear wordsmithed"


def grade_run(outputs_dir, assertions):
    """Grade a single run against its assertions."""
    results = []

    # Find .docx files
    docx_files = glob.glob(os.path.join(outputs_dir, '*.docx'))
    docx_text = ''
    docx_path = None
    if docx_files:
        docx_path = docx_files[0]
        docx_text = extract_docx_text(docx_path)

    for assertion in assertions:
        name = assertion['name']
        atype = assertion.get('type', 'custom')
        check = assertion.get('check', '')
        result = {"text": name, "passed": False, "evidence": ""}

        if atype == 'file_exists':
            if docx_files:
                result["passed"] = True
                result["evidence"] = f"Found: {os.path.basename(docx_files[0])}"
            else:
                result["evidence"] = f"No .docx files found in {outputs_dir}"

        elif atype == 'content_contains':
            if check in docx_text:
                result["passed"] = True
                result["evidence"] = f"'{check}' found in document"
            else:
                result["evidence"] = f"'{check}' NOT found in document text"

        elif atype == 'custom':
            if 'line_items' in name or 'line_item' in name:
                if docx_path:
                    count = count_line_items(docx_path)
                    expected = 2  # both eval 2 and 3 expect 2 items
                    result["passed"] = count == expected
                    result["evidence"] = f"Found {count} line items (expected {expected})"
                else:
                    result["evidence"] = "No .docx file to check"

            elif 'wordsmith' in name:
                if docx_path:
                    raw_inputs = ["march retainer", "data pipeline work"]
                    passed, evidence = check_wordsmithed(docx_path, raw_inputs)
                    result["passed"] = passed
                    result["evidence"] = evidence
                else:
                    result["evidence"] = "No .docx file to check"
            else:
                result["evidence"] = f"Custom assertion '{name}' — manual review needed"

        results.append(result)

    return results


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 grade_outputs.py <eval_dir> [variant]")
        print("  eval_dir: e.g., .../iteration-1/single-item-invoice")
        print("  variant: with_skill or without_skill (default: both)")
        sys.exit(1)

    eval_dir = Path(sys.argv[1])
    variants = [sys.argv[2]] if len(sys.argv) > 2 else ['with_skill', 'without_skill']

    # Load assertions
    meta_path = eval_dir / 'eval_metadata.json'
    with open(meta_path) as f:
        meta = json.load(f)
    assertions = meta.get('assertions', [])

    for variant in variants:
        outputs_dir = eval_dir / variant / 'outputs'
        if not outputs_dir.exists():
            print(f"Skipping {variant} — no outputs directory")
            continue

        results = grade_run(str(outputs_dir), assertions)
        passed = sum(1 for r in results if r['passed'])
        total = len(results)

        grading = {
            "eval_id": meta['eval_id'],
            "eval_name": meta['eval_name'],
            "variant": variant,
            "expectations": results,
            "pass_rate": f"{passed}/{total}"
        }

        grading_path = eval_dir / variant / 'grading.json'
        with open(grading_path, 'w') as f:
            json.dump(grading, f, indent=2)

        print(f"{meta['eval_name']} [{variant}]: {passed}/{total} passed")
        for r in results:
            status = "PASS" if r['passed'] else "FAIL"
            print(f"  [{status}] {r['text']}: {r['evidence']}")
        print()


if __name__ == '__main__':
    main()
