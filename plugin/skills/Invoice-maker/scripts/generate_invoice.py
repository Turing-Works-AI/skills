#!/usr/bin/env python3
"""
Turing Works Invoice Generator

Generates branded .docx invoices by modifying a template document.
Handles GST calculations, QR code generation, and multiple line items.

Usage:
    python3 generate_invoice.py <config.json> [output.docx]

Config JSON fields:
    invoice_code      : str   — e.g. "LYK0001" (3-letter client code + 4-digit number)
    client_name       : str   — e.g. "Lyka Pet Food Pty Ltd"
    client_address    : list  — e.g. ["Level 1, 123 Example St", "Sydney", "NSW", "2000"]
    line_items        : list  — each: {item, description, amount_inc_gst}
    payment_link      : str   — Wise payment URL
    bank_reference    : str   — bank transfer reference number
    invoice_date      : str   — optional, "YYYY-MM-DD", defaults to today
"""

import json
import sys
import os
import tempfile
from datetime import datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from copy import deepcopy

# ---------------------------------------------------------------------------
# Auto-install dependencies
# ---------------------------------------------------------------------------
def _ensure(module, package):
    try:
        __import__(module)
    except ImportError:
        os.system(f'{sys.executable} -m pip install "{package}" --quiet --break-system-packages 2>/dev/null || {sys.executable} -m pip install "{package}" --quiet')

_ensure('docx', 'python-docx')
_ensure('qrcode', 'qrcode[pil]')
_ensure('PIL', 'Pillow')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
import qrcode

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).parent
SKILL_DIR = SCRIPT_DIR.parent
TEMPLATE_PATH = SKILL_DIR / 'assets' / 'template.docx'

GST_RATE = Decimal('0.10')
PAYMENT_TERMS_DAYS = 14

BANK = {
    'account_name': 'Turing Works',
    'bsb': '774001',
    'account_number': '231410692',
}

WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def format_currency(amount):
    """$X,XXX.XX format."""
    return f'${amount:,.2f}'


def calc_gst(amount_inc_gst):
    """Return (ex_gst, gst) from an amount including GST at 10%."""
    inc = Decimal(str(amount_inc_gst))
    ex = (inc / (1 + GST_RATE)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    return ex, inc - ex


def generate_qr(url, output_path):
    """Generate a QR code PNG from a payment URL."""
    q = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_H,
                       box_size=10, border=2)
    q.add_data(url)
    q.make(fit=True)
    img = q.make_image(fill_color='#1B3A26', back_color='#F2F0EB')
    img.save(str(output_path))


def clear_cell(cell):
    """Remove all content from a table cell, preserving formatting."""
    # Keep only the first paragraph, clear everything inside it
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for child in list(p0._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink'):
            p0._element.remove(child)


def add_lines(cell, lines):
    """Write formatted lines into a cleared cell.

    lines: list of tuples (text, bold, size_pt, color).
    Each tuple becomes its own paragraph.
    """
    for i, (text, bold, size_pt, color) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size_pt:
            run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color


def insert_item_rows(table, extra_count):
    """Clone the first line-item row and insert copies after the second data row.

    The template has 2 data rows (indices 1-2). This adds extra_count rows
    so that the table can hold 2 + extra_count line items total.
    """
    if extra_count <= 0:
        return
    ref_tr = table.rows[2]._tr          # insert after last existing data row
    source_tr = table.rows[1]._tr       # clone formatting from first data row
    for _ in range(extra_count):
        new_tr = deepcopy(source_tr)
        # blank out text content
        for tc in new_tr.findall(qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                for r in p.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        t.text = ''
        ref_tr.addnext(new_tr)
        ref_tr = new_tr                 # next clone goes after this one


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_invoice(config, output_path):
    """Build an invoice .docx from a config dict and save to output_path."""

    # -- Parse config ----------------------------------------------------------
    invoice_code = config['invoice_code']           # e.g. "LYK0001"
    client_name  = config['client_name']
    client_addr  = config['client_address']         # list of address lines
    payment_link = config['payment_link']
    bank_ref     = config['bank_reference']
    items        = config['line_items']

    inv_date = datetime.strptime(config.get('invoice_date',
                    datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d')
    due_date = inv_date + timedelta(days=PAYMENT_TERMS_DAYS)

    # -- GST calculations -----------------------------------------------------
    total_inc = Decimal('0')
    for item in items:
        inc = Decimal(str(item['amount_inc_gst']))
        ex, gst = calc_gst(inc)
        item['_ex'], item['_gst'], item['_inc'] = ex, gst, inc
        total_inc += inc

    # -- Open template ---------------------------------------------------------
    doc = Document(str(TEMPLATE_PATH))
    t0 = doc.tables[0]   # header table
    t1 = doc.tables[1]   # line-items + footer table

    # =========================================================================
    # TABLE 0 — Header
    # =========================================================================

    # [1,1] Invoice title + entity
    cell = t0.cell(1, 1)
    clear_cell(cell)
    add_lines(cell, [
        ('', False, None, None),                         # spacing
        (f'Invoice #{invoice_code}', True, 36, WHITE),
        ('GENERAL PARTNERSHIP', False, 9, WHITE),
        ('ABN: 78 723 794 381', False, 9, WHITE),
    ])

    # [3,1] Bill-to
    cell = t0.cell(3, 1)
    clear_cell(cell)
    lines = [('BILL TO:', True, 8, None), (client_name, True, None, None)]
    for addr_line in client_addr:
        lines.append((addr_line, False, 9, None))
    add_lines(cell, lines)

    # [3,3] Invoice metadata
    cell = t0.cell(3, 3)
    clear_cell(cell)
    add_lines(cell, [
        ('INVOICE #',        True,  8,    None),
        (invoice_code,       False, 9,    None),
        ('DATE',             True,  8,    None),
        (inv_date.strftime('%d/%m/%y'), False, 9, None),
        ('INVOICE DUE DATE', True,  8,    None),
        (due_date.strftime('%d/%m/%y'), False, 9, None),
    ])

    # =========================================================================
    # TABLE 1 — Line items + footer
    # =========================================================================

    # Insert extra rows if more than 2 line items
    if len(items) > 2:
        insert_item_rows(t1, len(items) - 2)

    # Fill line-item rows (indices 1 … len(items))
    for idx, item in enumerate(items):
        row = t1.rows[1 + idx]
        data = {
            1: item['item'],
            2: item['description'],
            3: format_currency(item['_ex']),
            4: '10%',
            5: format_currency(item['_gst']),
            6: format_currency(item['_inc']),
        }
        for ci, text in data.items():
            c = row.cells[ci]
            clear_cell(c)
            add_lines(c, [(text, False, 9, None)])

    # Clear any remaining empty data rows (between last item and footer)
    first_empty = 1 + len(items)
    # Footer starts 4 rows from the end: spacer, label, data, payment
    footer_start = len(t1.rows) - 4
    for ri in range(first_empty, footer_start):
        for ci in range(1, 7):
            c = t1.rows[ri].cells[ci]
            clear_cell(c)

    # -- Footer: bank details --------------------------------------------------
    # Bank data row is 2nd-to-last row
    bank_data_ri = len(t1.rows) - 2
    cell = t1.rows[bank_data_ri].cells[1]   # merged cols 1-2
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    bank_lines = [
        ('Account Name   ', BANK['account_name']),
        ('BSB code           ', BANK['bsb']),
        ('Acc/No                ', BANK['account_number']),
        ('Reference           ', bank_ref),
    ]
    for i, (label, value) in enumerate(bank_lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_val = p.add_run(value)
        run_val.font.size = Pt(10)

    # -- Footer: total amount --------------------------------------------------
    cell = t1.rows[bank_data_ri].cells[3]   # merged cols 3-6
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(format_currency(total_inc))
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = WHITE

    # -- Replace QR code image -------------------------------------------------
    qr_tmp = Path(tempfile.mktemp(suffix='.png'))
    generate_qr(payment_link, qr_tmp)
    try:
        for rel in doc.part.rels.values():
            if 'image' in rel.reltype and 'image2' in str(rel.target_ref):
                with open(qr_tmp, 'rb') as f:
                    rel.target_part._blob = f.read()
                break
    finally:
        qr_tmp.unlink(missing_ok=True)

    # -- Update payment link paragraph -----------------------------------------
    for p in doc.paragraphs:
        if p.text.strip().startswith('http'):
            for child in list(p._element):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('r', 'hyperlink'):
                    p._element.remove(child)
            run = p.add_run(payment_link)
            run.italic = True
            run.font.size = Pt(8)
            break

    # -- Save ------------------------------------------------------------------
    doc.save(str(output_path))
    print(f'Invoice generated: {output_path}')
    return str(output_path)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print('Usage: python3 generate_invoice.py <config.json> [output.docx]')
        sys.exit(1)

    with open(sys.argv[1]) as f:
        config = json.load(f)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        inv_date = datetime.strptime(
            config.get('invoice_date', datetime.now().strftime('%Y-%m-%d')),
            '%Y-%m-%d')
        output_path = f'{inv_date.strftime("%Y%m%d")}_Invoice_{config["invoice_code"]}.docx'

    generate_invoice(config, output_path)


if __name__ == '__main__':
    main()
